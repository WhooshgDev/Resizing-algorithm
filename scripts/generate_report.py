"""Generate the 10-15 page Vietnamese report as .docx"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ─── Styles ────────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(13)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Times New Roman'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    hs.font.color.rgb = RGBColor(0, 0, 0)

# ─── Helper ─────────────────────────────────────────────────────────────────
def add_para(text, bold=False, italic=False, size=None, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size: run.font.size = Pt(size)
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5 + level * 1)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(50, 50, 50)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1)
    return p

# ═══════════════════════════════════════════════════════════════════════════
# TRANG BIA
# ═══════════════════════════════════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

add_para("BỘ GIÁO DỤC VÀ ĐÀO TẠO", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("TRƯỜNG ĐẠI HỌC ...", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("KHOA ...", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_para("---o0o---", bold=False, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_para("BÁO CÁO ĐỒ ÁN", bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("XỬ LÝ ẢNH SỐ", bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("CHỦ ĐỀ: THUẬT TOÁN SEAM CARVING", bold=True, size=15, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("VÀ CÁC PHƯƠNG PHÁP THAY ĐỔI KÍCH THƯỚC ẢNH", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
doc.add_paragraph()
add_para("Giảng viên hướng dẫn: ...", size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Sinh viên thực hiện: ...", size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Mã số sinh viên: ...", size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Lớp: ...", size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_para("TP. Hồ Chí Minh, tháng 6 năm 2026", size=13, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# MUC LUC
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('MỤC LỤC', level=1)
toc_items = [
    ("1.", "Mô tả bài toán", 3),
    ("2.", "Lý thuyết và kiến thức liên quan", 4),
    ("", "2.1. Các phép biến đổi ảnh cơ bản", 4),
    ("", "2.2. Thuật toán Seam Carving", 5),
    ("", "2.3. Forward Energy", 6),
    ("", "2.4. Multi-Operator Retargeting", 7),
    ("", "2.5. Phân loại ảnh tự động", 7),
    ("3.", "Cách giải quyết", 8),
    ("", "3.1. Kiến trúc chương trình", 8),
    ("", "3.2. Các thuật toán cài đặt", 9),
    ("", "3.3. Adaptive Hybrid", 10),
    ("", "3.4. Auto-classifier", 10),
    ("4.", "Kết quả thực nghiệm", 11),
    ("5.", "Bài học và trải nghiệm", 13),
    ("6.", "Tài liệu tham khảo", 14),
]
for num, title, page in toc_items:
    indent = "     " if num == "" else " "
    txt = f"{num}{indent}{title}  {'·' * (50 - len(num+title))}  {page}"
    add_para(txt, size=12, space_after=2)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# CHUONG 1: MO TA BAI TOAN
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('1. Mô tả bài toán', level=1)

add_para("Trong xử lý ảnh số, bài toán thay đổi kích thước ảnh (image resizing) là một trong những bài toán "
         "cơ bản nhưng quan trọng. Mục tiêu là biến đổi một ảnh gốc có kích thước W1 × H1 thành ảnh mới "
         "có kích thước W2 × H2 sao cho nội dung quan trọng của ảnh được bảo toàn tối đa.", size=13)

add_para("Các phương pháp truyền thống bao gồm:", size=13)
add_bullet("Scaling (co giãn): Sử dụng nội suy để thay đổi kích thước. Nhược điểm: làm méo mó tỉ lệ, "
           "mất chi tiết khi phóng to/thu nhỏ quá mức.")
add_bullet("Cropping (cắt xén): Giữ lại một vùng của ảnh. Nhược điểm: mất thông tin ở vùng bị cắt.")
add_bullet("Letterboxing: Co giãn giữ nguyên tỉ lệ, thêm viền đen. Nhược điểm: lãng phí diện tích hiển thị.")

add_para("Năm 2007, Avidan và Shamir giới thiệu thuật toán Seam Carving (content-aware image resizing), "
         "một cách tiếp cận hoàn toàn mới: thay vì co giãn toàn bộ ảnh, thuật toán tìm và loại bỏ các đường "
         "nối (seam) có năng lượng thấp nhất, giúp bảo toàn các vùng nội dung quan trọng.", size=13)

add_para("Bài toán đặt ra trong đồ án này là xây dựng một hệ thống thông minh có khả năng:", size=13)
add_bullet("Phân loại tự động loại ảnh (ảnh chụp, đồ họa, văn bản...)")
add_bullet("Áp dụng thuật toán phù hợp nhất cho từng loại ảnh")
add_bullet("Cài đặt thuật toán Seam Carving với Forward Energy (cải tiến chất lượng)")
add_bullet("Xây dựng thuật toán lai (hybrid) kết hợp seam carving và scaling")
add_bullet("Tạo giao diện so sánh trực quan giữa các phương pháp")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# CHUONG 2: LY THUYET
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('2. Lý thuyết và kiến thức liên quan', level=1)

doc.add_heading('2.1. Các phép biến đổi ảnh cơ bản', level=2)

add_para("a) Scaling (co giãn ảnh)", bold=True, size=13)
add_para("Scaling là phép biến đổi几何 (geometric) cơ bản nhất. Ảnh gốc được ánh xạ sang lưới điểm "
         "mới thông qua các phương pháp nội suy (interpolation):", size=13)
add_bullet("Nearest-neighbor: Nhanh nhất nhưng chất lượng thấp, gây hiệu ứng răng cưa.")
add_bullet("Bilinear: Nội suy tuyến tính trên 2 chiều, chất lượng trung bình.")
add_bullet("Bicubic: Nội suy bậc 3 trên lân cận 4×4, chất lượng tốt.")
p = doc.add_paragraph(style='List Bullet')
run = p.add_run("Lanczos: Sử dụng hàm sinc làm nhân nội suy (Lanczos-3 truncation), "
           "cho chất lượng cao nhất trong các phương pháp phổ biến. Đây là phương pháp được "
           "sử dụng trong đồ án này cho phép scaling.")
run.italic = True
p.paragraph_format.space_after = Pt(2)

add_para("b) Cropping (cắt xén)", bold=True, size=13)
add_para("Cropping giữ lại một vùng hình chữ nhật của ảnh gốc và loại bỏ phần còn lại. "
         "Crop thông minh (smart crop) sử dụng entropy để xác định vùng có nhiều thông tin nhất.", size=13)

add_para("c) Letterboxing", bold=True, size=13)
add_para("Kỹ thuật co giãn ảnh giữ nguyên tỉ lệ khung hình gốc, sau đó thêm viền màu (thường là đen) "
         "xung quanh để lấp đầy kích thước mục tiêu.", size=13)

doc.add_heading('2.2. Thuật toán Seam Carving', level=2)

add_para("Seam Carving (Avidan & Shamir, SIGGRAPH 2007) là thuật toán thay đổi kích thước ảnh có nhận "
         "thức nội dung (content-aware image resizing). Ý tưởng chính: thay vì co giãn toàn bộ ảnh, "
         "ta tìm và loại bỏ các đường nối (seam) có năng lượng thấp nhất.", size=13)

add_para("Định nghĩa: Một seam dọc là đường đi từ trên xuống dưới, mỗi hàng chọn đúng một pixel, "
         "và độ lệch cột giữa hai hàng liên tiếp không quá 1:", size=13)
add_para("S^x = {s_i^x}_{i=1}^n = {x(i)}_i=1^n, với |x(i) - x(i-1)| ≤ 1", italic=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Năng lượng của seam là tổng năng lượng của các pixel trên đường đi đó:", size=13)
add_para("E(S) = Σ_{i=1}^n e(I(x(i), i))", italic=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

add_para("Hàm năng lượng e(I) thường dùng là gradient magnitude (backward energy):", size=13)
add_para("e(I) = |∂I/∂x| + |∂I/∂y|", italic=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Tuy nhiên, backward energy chỉ đánh giá năng lượng tại pixel bị loại bỏ, không tính đến "
         "năng lượng mới được tạo ra khi các pixel lân cận trở thành kề nhau.", size=13)

add_para("Quy hoạch động (Dynamic Programming):", bold=True, size=13)
add_para("Để tìm seam có năng lượng thấp nhất, sử dụng quy hoạch động với độ phức tạp O(n×m):", size=13)
add_code("M[i,j] = e(i,j) + min(M[i-1,j-1], M[i-1,j], M[i-1,j+1])")
add_para("Sau đó truy vết (backtracking) từ hàng cuối cùng để tìm đường đi tối ưu.", size=13)

doc.add_heading('2.3. Forward Energy', level=2)

add_para("Forward Energy (Rubinstein et al., SIGGRAPH 2008) là cải tiến quan trọng của Seam Carving. "
         "Thay vì chỉ dựa vào năng lượng tại pixel bị xóa, forward energy tính toán năng lượng "
         "của các cạnh mới được tạo ra giữa các pixel lân cận sau khi xóa seam.", size=13)

add_para("Công thức forward energy:", bold=True, size=13)
add_para("Với mỗi pixel (i,j), có 3 trường hợp dựa trên hướng seam đi vào:", size=13)
add_bullet("Seam từ trên-trái (x(i) = x(i-1) - 1):")
add_code("CL(i,j) = |I(i,j+1)-I(i,j-1)| + |I(i-1,j)-I(i,j-1)|")
add_bullet("Seam từ trên-xuống thẳng (x(i) = x(i-1)):")
add_code("CU(i,j) = |I(i,j+1)-I(i,j-1)|")
add_bullet("Seam từ trên-phải (x(i) = x(i-1) + 1):")
add_code("CR(i,j) = |I(i,j+1)-I(i,j-1)| + |I(i-1,j)-I(i,j+1)|")

add_para("Công thức DP với forward energy:", size=13)
add_code("dp[i,j] = min(dp[i-1,j-1]+CL[i,j], dp[i-1,j]+CU[i,j], dp[i-1,j+1]+CR[i,j])")

add_para("Forward energy cho chất lượng ảnh tốt hơn đáng kể so với backward energy, đặc biệt "
         "khi loại bỏ nhiều seam trên cùng một ảnh.", size=13)

doc.add_heading('2.4. Multi-Operator Retargeting', level=2)

add_para("Multi-Operator Media Retargeting (Rubinstein et al., SIGGRAPH 2009) là phương pháp kết hợp "
         "nhiều toán tử (seam carving, scaling, cropping) để thay đổi kích thước ảnh. Ý tưởng chính: "
         "mỗi toán tử có ưu điểm riêng và phù hợp với từng loại nội dung ảnh khác nhau.", size=13)

add_para("Phương pháp lai (Hybrid) được đề xuất:", bold=True, size=13)
add_bullet("Giai đoạn 1: Áp dụng seam carving để loại bỏ các vùng có năng lượng thấp (nền đồng nhất, "
           "bầu trời, tường...).")
add_bullet("Giai đoạn 2: Khi năng lượng seam bắt đầu tăng cao (chạm vào vùng nội dung quan trọng), "
           "chuyển sang scaling.")
add_bullet("Điểm then chốt: xác định thời điểm chuyển đổi giữa hai giai đoạn dựa trên ngưỡng "
           "năng lượng động (adaptive threshold).")

doc.add_heading('2.5. Phân loại ảnh tự động', level=2)

add_para("Để tự động chọn thuật toán phù hợp, hệ thống cần phân loại ảnh đầu vào. Các đặc trưng "
         "được sử dụng:", size=13)

add_para("a) Mật độ cạnh (Edge density):", bold=True, size=13)
add_para("Tỉ lệ pixel có năng lượng > 50% năng lượng tối đa. Ảnh văn bản có mật độ cạnh cao "
         "(nhiều đường nét rõ). Ảnh chụp có mật độ cạnh thấp hơn.", size=13)

add_para("b) Độ phong phú màu sắc (Color richness):", bold=True, size=13)
add_para("Tỉ lệ màu duy nhất trên tổng số pixel (lấy mẫu ngẫu nhiên 4096 pixel). "
         "Ảnh đồ họa có ít màu, ảnh chụp có nhiều màu.", size=13)

add_para("c) Entropy năng lượng (Energy entropy):", bold=True, size=13)
add_para("Entropy chuẩn hóa của phân bố năng lượng. Đo lường mức độ phân tán của gradient trong ảnh. "
         "Ảnh chụp có entropy cao (nhiều vùng khác nhau).", size=13)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# CHUONG 3: CACH GIAI QUYET
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('3. Cách giải quyết', level=1)

doc.add_heading('3.1. Kiến trúc chương trình', level=2)

add_para("Chương trình được viết bằng Python 3, sử dụng thư viện NumPy cho tính toán số học "
         "và Pillow (PIL) cho đọc/ghi ảnh. Kiến trúc gồm 4 module chính:", size=13)

add_para("Module 1: Energy & Seam (compute_energy, _find_seam_forward, _find_seam_backward)", bold=True)
add_para("Tính năng lượng ảnh bằng gradient magnitude (backward energy) và forward energy cost. "
         "Tìm seam tối ưu bằng quy hoạch động.", size=13)

add_para("Module 2: Carving drivers (_carve_seams, _carve_adaptive)", bold=True)
add_para("Điều phối quá trình loại bỏ seam: loại bỏ N seam tuần tự, hoặc loại bỏ thích ứng "
         "dừng lại khi năng lượng seam vượt ngưỡng.", size=13)

add_para("Module 3: Algorithm implementations (resize_seam, resize_scale, resize_crop, resize_letterbox, resize_hybrid)", bold=True)
add_para("Cài đặt 5 thuật toán thay đổi kích thước ảnh.", size=13)

add_para("Module 4: Classifier & Dispatcher (classify_image, resize_auto)", bold=True)
add_para("Phân loại ảnh và tự động chọn thuật toán phù hợp.", size=13)

doc.add_heading('3.2. Các thuật toán cài đặt', level=2)

add_para("a) Seam Carving:", bold=True, size=13)
add_para("Sử dụng forward energy với quy hoạch động O(h×w) cho mỗi seam. "
         "Hỗ trợ cả seam dọc (giảm chiều rộng) và seam ngang (giảm chiều cao). "
         "Khi giảm chiều cao, ảnh được chuyển vị, áp dụng seam carving dọc, "
         "sau đó chuyển vị lại.", size=13)

add_para("b) Standard Scaling:", bold=True, size=13)
add_para("Sử dụng Lanczos interpolation (Image.LANCZOS trong Pillow) cho chất lượng cao nhất.", size=13)

add_para("c) Smart Crop:", bold=True, size=13)
add_para("Chia ảnh thành các tile 16×16. Tính entropy cho mỗi tile. Sử dụng summed-area table "
         "(integral image) để tìm cửa sổ crop có tổng entropy lớn nhất trong O(m×n). "
         "Sau crop, scale về kích thước mục tiêu.", size=13)

add_para("d) Letterbox:", bold=True, size=13)
add_para("Scale ảnh giữ nguyên tỉ lệ khung hình, đặt vào giữa khung mục tiêu, "
         "tô viền bằng màu đen.", size=13)

add_para("e) Adaptive Hybrid:", bold=True, size=13)
add_para("Phương pháp lai thích ứng — kết hợp seam carving và scaling. "
         "Đây là đóng góp chính của đồ án, được mô tả chi tiết trong mục 3.3.", size=13)

doc.add_heading('3.3. Adaptive Hybrid (đóng góp chính)', level=2)

add_para("Adaptive Hybrid là cải tiến của Multi-Operator Retargeting, sử dụng ngưỡng năng lượng "
         "động để quyết định thời điểm chuyển từ seam carving sang scaling.", size=13)

add_para("Thuật toán chi tiết:", bold=True, size=13)
add_para("1. Khởi tạo ảnh làm việc và danh sách theo dõi năng lượng seam.", size=13)
add_para("2. Lặp tối đa max_n lần:", size=13)
add_bullet("Tìm seam có forward energy thấp nhất.")
add_bullet("Tính năng lượng chuẩn hóa của seam (fe / chiều cao ảnh).")
add_bullet("Cập nhật danh sách năng lượng seam.")
add_bullet("Kiểm tra điều kiện dừng:")
add_code("nonzero = [e for e in seam_energies if e > 1e-6]")
add_code("if len(nonzero) >= 5 and i >= 10:")
add_code("    baseline = median(nonzero)")
add_code("    if norm_e > baseline * energy_ratio: break")
add_bullet("Nếu không dừng: loại bỏ seam khỏi ảnh.")
add_para("3. Scale phần còn lại bằng Lanczos interpolation.", size=13)

add_para("Ba điểm then chốt của thuật toán:", bold=True, size=13)
add_bullet("Median thay vì mean: Median bền vững với nhiễu, không bị ảnh hưởng bởi một seam có năng lượng cao bất thường.")
add_bullet("Bỏ qua seam năng lượng zero: Các vùng đồng nhất (viền trắng, nền đen) có năng lượng "
           "forward ≈ 0. Nếu tính vào baseline, mọi seam đầu tiên chạm vào vùng có nội dung sẽ "
           "kích hoạt dừng sớm.")
add_bullet("Chỉ kiểm tra sau 10 seam đầu: Đảm bảo có đủ dữ liệu để tính median tin cậy.")

add_para("Tham số energy_ratio (mặc định 2.0) điều chỉnh độ 'bảo thủ':", size=13)
add_bullet("1.5 → Dừng sớm, scaling nhiều hơn, an toàn cho văn bản.")
add_bullet("2.0 → Cân bằng (được sử dụng trong tất cả thử nghiệm).")
add_bullet("3.0 → Carving nhiều hơn, tốt cho ảnh có nhiều vùng đồng nhất.")

doc.add_heading('3.4. Auto-classifier', level=2)

add_para("Bộ phân loại ảnh tự động sử dụng 3 đặc trưng để xác định loại ảnh:", size=13)

add_para("1. Mật độ cạnh (Edge density):", bold=True)
add_para("Tính trên ảnh thu nhỏ 256×256. Tỉ lệ pixel có gradient > 50% max gradient.", size=13)

add_para("2. Entropy năng lượng (Energy entropy):", bold=True)
add_para("Histogram năng lượng với 32 bins, entropy chuẩn hóa về [0,1]. "
         "Ảnh chụp tự nhiên có entropy cao.", size=13)

add_para("3. Độ phong phú màu (Color ratio):", bold=True)
add_para("Lấy mẫu ngẫu nhiên 4096 pixel. Đếm số màu duy nhất (mỗi màu là tổ hợp R,G,B 24-bit).", size=13)

add_para("Luật phân loại:", bold=True, size=13)
add_bullet("TEXT: edge_density > 0.35 và color_ratio < 0.15")
add_bullet("GRAPHIC: color_ratio < 0.10, hoặc (color_ratio < 0.25 và edge_density < 0.20)")
add_bullet("PHOTO: color_ratio > 0.30 và energy_entropy > 0.75")
add_bullet("DEFAULT: tất cả các trường hợp còn lại")

add_para("Sau khi phân loại, bộ điều phối chọn thuật toán:", size=13)
add_bullet("PHOTO, DEFAULT → hybrid")
add_bullet("GRAPHIC, TEXT → scale")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# CHUONG 4: KET QUA
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('4. Kết quả thực nghiệm', level=1)

add_para("Chương trình được thử nghiệm trên 24 ảnh (20 ảnh từ picsum.photos + 4 ảnh tự tạo: "
         "văn bản, đồ họa, biểu đồ, UI). Kết quả phân loại:", size=13)

# Classification results table
table = doc.add_table(rows=6, cols=6)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Loại ảnh', 'Phân loại', 'Edge density', 'Energy entr.', 'Color ratio', 'Thuật toán']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10)

data = [
    ['Photo (20 ảnh)', 'default/photo', '0.001-0.049', '0.073-0.174', '0.498-0.947', 'hybrid'],
    ['Text', 'graphic', '0.168', '0.110', '0.056', 'scale'],
    ['Chart', 'graphic', '0.047', '0.068', '0.072', 'scale'],
    ['UI', 'graphic', '0.041', '0.090', '0.080', 'scale'],
    ['Graphic', 'default', '0.029', '0.078', '0.498', 'hybrid'],
]
for row_idx, row_data in enumerate(data):
    for col_idx, val in enumerate(row_data):
        cell = table.rows[row_idx + 1].cells[col_idx]
        cell.text = val
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)

doc.add_paragraph()

add_para("Kết quả phân loại đúng 22/24 ảnh (91.7%). Ảnh đồ họa trừu tượng (các vòng tròn màu "
         "ngẫu nhiên) bị phân loại nhầm thành default → hybrid do độ phong phú màu cao.", size=13)

add_para("Kết quả Adaptive Hybrid so với Fixed 70%:", bold=True, size=13)

table2 = doc.add_table(rows=6, cols=4)
table2.style = 'Light Grid Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

headers2 = ['Loại ảnh', 'Seam carved', 'Scaled', 'Ghi chú']
for i, h in enumerate(headers2):
    cell = table2.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10)

data2 = [
    ['Photo', '400 + 15', '0 + 285', 'Bầu trời/nền bị carve, nội dung được scale'],
    ['Text', '96 + 102', '304 + 198', 'Lề trắng bị carve, chữ được scale'],
    ['Chart', '234 + 300', '166 + 0', 'Viền trắng bị carve, biểu đồ được giữ'],
    ['UI', '144 + 210', '256 + 90', 'Thanh công cụ bị carve, nội dung scale'],
    ['Graphic', '365 + 300', '35 + 0', 'Vùng thưa bị carve, vùng dày được giữ'],
]
for row_idx, row_data in enumerate(data2):
    for col_idx, val in enumerate(row_data):
        cell = table2.rows[row_idx + 1].cells[col_idx]
        cell.text = val
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)

doc.add_paragraph()

add_para("Nhận xét:", bold=True, size=13)
add_bullet("Với ảnh chụp: Adaptive Hybrid carve gần như toàn bộ phần nền đồng nhất (bầu trời, mặt đất) "
           "và dừng lại khi chạm vào cây cối/tòa nhà.")
add_bullet("Với văn bản: Hệ thống carve các lề trắng (≈96 seam dọc + 102 seam ngang) rồi chuyển "
           "sang scale, tránh làm méo chữ.")
add_bullet("Với biểu đồ: Tất cả viền trắng bị carve, biểu đồ được giữ nguyên vẹn nhờ scaling.")
add_bullet("Thời gian xử lý: ảnh 800×600 → 400×300 mất 10-60 giây tùy loại ảnh và thuật toán.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# CHUONG 5: BAI HOC
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('5. Bài học và trải nghiệm thu được', level=1)

add_para("Qua quá trình thực hiện đồ án này, chúng em đã học được nhiều kiến thức và kỹ năng bổ ích:", size=13)

add_para("5.1. Kiến thức về xử lý ảnh số", bold=True, size=13)
add_bullet("Hiểu sâu về các phép biến đổi ảnh cơ bản: scaling, cropping, interpolation.")
add_bullet("Nắm vững thuật toán Seam Carving và các cải tiến (Forward Energy, Multi-Operator Retargeting).")
add_bullet("Biết cách đánh giá chất lượng ảnh sau khi xử lý thông qua các metric như năng lượng, entropy.")

add_para("5.2. Kỹ năng lập trình", bold=True, size=13)
add_bullet("Sử dụng thành thạo NumPy cho tính toán số học trên ma trận lớn (vectorization, broadcasting).")
add_bullet("Kỹ thuật tối ưu hóa code Python: tránh vòng lặp, sử dụng boolean mask, rolling array.")
add_bullet("Xây dựng kiến trúc phần mềm module hóa, dễ bảo trì và mở rộng.")

add_para("5.3. Kỹ năng giải quyết vấn đề", bold=True, size=13)
add_bullet("Phát hiện và sửa lỗi trong quá trình cài đặt thuật toán (lỗi broadcast, lỗi chỉ số mảng, "
           "lỗi logíc DP).")
add_bullet("Điều chỉnh tham số (energy_ratio) để cân bằng giữa chất lượng và tốc độ.")
add_bullet("Xây dựng bộ phân loại ảnh tự động dựa trên đặc trưng thống kê.")

add_para("5.4. Bài học về thiết kế thuật toán", bold=True, size=13)
add_bullet("Không có thuật toán nào 'tốt nhất' cho mọi loại ảnh. Mỗi thuật toán có ưu nhược điểm riêng.")
add_bullet("Cách tiếp cận hybrid (kết hợp nhiều phương pháp) thường cho kết quả tốt hơn từng phương pháp riêng lẻ.")
add_bullet("Adaptive threshold dựa trên median bền vững hơn mean khi có nhiễu.")
add_bullet("Việc bỏ qua các giá trị zero trước khi tính baseline giúp tránh dừng sớm không mong muốn.")

add_para("5.5. Hướng phát triển", bold=True, size=13)
add_bullet("Tích hợp phát hiện khuôn mặt (face detection) để bảo vệ khuôn mặt khi seam carving.")
add_bullet("Sử dụng deep learning để phân loại ảnh chính xác hơn (thay vì dùng đặc trưng thống kê đơn giản).")
add_bullet("Song song hóa (multi-threading) quá trình seam carving để tăng tốc độ.")
add_bullet("Hỗ trợ thêm các định dạng ảnh và giao diện người dùng đồ họa (GUI).")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# TAI LIEU THAM KHAO
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('6. Tài liệu tham khảo', level=1)

refs = [
    "[1] Avidan, S., & Shamir, A. (2007). Seam carving for content-aware image resizing. ACM Transactions on Graphics (TOG), 26(3), 10-es.",
    "[2] Rubinstein, M., Shamir, A., & Avidan, S. (2008). Improved seam carving for video retargeting. ACM Transactions on Graphics (TOG), 27(3), 1-9.",
    "[3] Rubinstein, M., Shamir, A., & Avidan, S. (2009). Multi-operator media retargeting. ACM Transactions on Graphics (TOG), 28(3), 1-11.",
    "[4] Viola, P., & Jones, M. (2001). Rapid object detection using a boosted cascade of simple features. CVPR 2001.",
    "[5] González, R. C., & Woods, R. E. (2018). Digital Image Processing (4th ed.). Pearson.",
    "[6] numpy.org. NumPy Documentation. https://numpy.org/doc/",
    "[7] pillow.readthedocs.io. Pillow (PIL Fork) Documentation. https://pillow.readthedocs.io/",
    "[8] python-docx.readthedocs.io. python-docx Documentation. https://python-docx.readthedocs.io/",
]
for ref in refs:
    add_para(ref, size=11, space_after=4)

# ─── Save ───────────────────────────────────────────────────────────────────
path = "BaoCao_DoAn_XuLyAnhSo_SeamCarving.docx"
doc.save(path)
print(f"Saved: {path} ({os.path.getsize(path)/1024:.0f} KB)")
